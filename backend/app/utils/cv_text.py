"""
Text extraction utilities for CV/resume processing
"""

import os
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from various file formats (PDF, DOCX, TXT)
    """
    p = Path(file_path)
    if not p.exists() or p.stat().st_size == 0:
        raise ValueError("Cannot read an empty file")
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_lower = file_path.lower()

    if file_lower.endswith(".pdf"):
        return _extract_pdf_text(file_path)
    if file_lower.endswith(".docx"):
        return _extract_docx_text(file_path)
    if file_lower.endswith(".txt"):
        return _extract_txt_text(file_path)
    raise ValueError(
        f"Unsupported file type: {os.path.basename(file_path)}. Supported types: PDF, DOCX, TXT"
    )


def _extract_pdf_text(file_path: str) -> str:
    reader = PdfReader(file_path)
    parts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t)
    return "\n".join(parts)


def _extract_docx_text(file_path: str) -> str:
    doc = Document(file_path)
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_txt_text(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# Heuristic: collapse sequences of single letters like "R e a c t" -> "React"
# Keeps numbers/emails/URLs intact.
_SPACED_LETTERS_RE = re.compile(
    r"(?:\b[a-zA-Z]\s){4,}[a-zA-Z]\b"
)  # ≥5 letters with spaces


def _fix_spaced_letters(text: str) -> str:
    def _join_run(run: str) -> str:
        # run like "R e a c t" -> "React"
        letters = run.split()
        return "".join(letters)

    out = []
    i = 0
    while i < len(text):
        m = _SPACED_LETTERS_RE.search(text, i)
        if not m:
            out.append(text[i:])
            break
        # append chunk before match
        out.append(text[i : m.start()])
        # replace the matched run
        out.append(_join_run(m.group(0)))
        i = m.end()
    return "".join(out)


def clean_extracted_text(text: str) -> str:
    """
    Clean and normalize extracted text
    """
    if not text:
        return ""
    # collapse whitespace
    text = " ".join(text.split())
    # remove non-printable (keep basic punctuation)
    text = re.sub(r"[^\w\s.,!?-]", " ", text)
    # normalize spaces again
    text = " ".join(text.split())
    # NEW: de-space letter-by-letter artifacts
    text = _fix_spaced_letters(text)
    return text.strip()


def get_file_info(file_path: str) -> dict:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    stat = os.stat(file_path)
    _, ext = os.path.splitext(file_path)
    return {
        "path": file_path,
        "filename": os.path.basename(file_path),
        "extension": ext.lower(),
        "size_bytes": stat.st_size,
        "size_kb": round(stat.st_size / 1024, 2),
        "modified_time": stat.st_mtime,
    }


def preview_text(text: str, max_length: int = 500) -> str:
    if not text:
        return ""
    if len(text) <= max_length:
        return text
    return text[:max_length].strip() + "..."
